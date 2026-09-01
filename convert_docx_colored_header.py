import os
import sys
import json
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

def rgb_to_hex(rgb):
    if not rgb:
        return None
    try:
        return f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
    except Exception:
        return None

def get_run_style(run, is_header=False):
    styles = []
    font = run.font

    font_name = font.name if font.name else 'Noto Sans Devanagari'
    styles.append(f"font-family: '{font_name}', 'Noto Sans Devanagari', 'Mangal', sans-serif")

    if font.size and hasattr(font.size, 'pt') and font.size.pt:
        styles.append(f"font-size: {font.size.pt}pt")

    if font.color and font.color.rgb:
        hex_color = rgb_to_hex(font.color.rgb)
        if hex_color and hex_color.upper() == '#D9E8F5':
            hex_color = '#1E40AF'
        if hex_color:
            styles.append(f"color: {hex_color}")
    elif is_header:
        styles.append("color: #ffffff")

    if font.bold is True:
        styles.append("font-weight: bold")
    else:
        styles.append("font-weight: normal")

    if font.italic is True:
        styles.append("font-style: italic")
    else:
        styles.append("font-style: normal")

    if font.underline is True:
        styles.append("text-decoration: underline")

    return "; ".join(styles)

def get_paragraph_style(p, is_header=False, is_footer=False):
    styles = ["font-style: normal"]
    pf = p.paragraph_format

    if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        styles.append("text-align: center")
    elif p.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        styles.append("text-align: right")
    elif p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        styles.append("text-align: justify")
    elif p.alignment == WD_ALIGN_PARAGRAPH.LEFT:
        styles.append("text-align: left")
    elif is_header:
        styles.append("text-align: center")
    elif is_footer:
        styles.append("text-align: right")

    if pf.line_spacing:
        if isinstance(pf.line_spacing, (int, float)):
            styles.append(f"line-height: {pf.line_spacing}")
        elif hasattr(pf.line_spacing, 'pt') and pf.line_spacing.pt:
            styles.append(f"line-height: {pf.line_spacing.pt}pt")
    else:
        styles.append("line-height: 1.6")

    if pf.space_before and hasattr(pf.space_before, 'pt') and pf.space_before.pt:
        styles.append(f"margin-top: {pf.space_before.pt}pt")
    if pf.space_after and hasattr(pf.space_after, 'pt') and pf.space_after.pt:
        styles.append(f"margin-bottom: {pf.space_after.pt}pt")
    else:
        styles.append("margin-bottom: 0.4rem")

    return "; ".join(styles)

def paragraph_to_html(p, is_header=False, is_footer=False):
    p_style = get_paragraph_style(p, is_header=is_header, is_footer=is_footer)
    runs_html = []

    for run in p.runs:
        text = run.text
        if not text:
            continue
        text_escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
        r_style = get_run_style(run, is_header=is_header)
        runs_html.append(f'<span style="{r_style}">{text_escaped}</span>')

    full_inner = "".join(runs_html)
    if not full_inner.strip():
        return '<p style="margin-bottom: 0.2rem;"><br/></p>'

    return f'<p style="{p_style}; font-style: normal;">{full_inner}</p>'

def table_to_html(table):
    rows_html = []
    for row in table.rows:
        cells_html = []
        for cell in row.cells:
            cell_paragraphs = [paragraph_to_html(p) for p in cell.paragraphs]
            cell_content = "".join(cell_paragraphs)
            cells_html.append(f'<td style="border: 1px solid #CBD5E1; padding: 10px 14px; vertical-align: top; font-style: normal;">{cell_content}</td>')
        rows_html.append(f'<tr>{"".join(cells_html)}</tr>')
    return f'<table style="width: 100%; border-collapse: collapse; margin: 1.25rem 0; border: 1px solid #CBD5E1; font-style: normal;"><tbody>{"".join(rows_html)}</tbody></table>'

CAT_TITLES = {
    'summary': 'पाठ सारांश (Summary)',
    'notes': 'प्रश्न-उत्तर एवं व्याख्या (Notes)',
    'competency': 'योग्यता आधारित प्रश्न (Competency Based Questions)',
    'additional': 'अतिरिक्त प्रश्न (Additional Questions)',
    'muhavre': 'शब्द-अर्थ एवं मुहावरे (Word Meanings & Idioms)'
}

def docx_to_exact_html(filepath, cat_key='summary'):
    doc = docx.Document(filepath)
    html_parts = []

    cat_title = CAT_TITLES.get(cat_key, '')

    # Standardized Colored Banner Header (18pt title size across all content)
    header_html = f'''<div class="doc-header" style="background: linear-gradient(135deg, #156082 0%, #0F4C6A 100%); padding: 1.35rem 1.5rem; border-radius: 14px; margin-bottom: 2rem; text-align: center; color: #ffffff; box-shadow: 0 6px 18px rgba(21,96,130,0.22); font-style: normal;">
  <div class="doc-header-ch-name" style="font-size: 18pt; font-weight: bold; color: #ffffff !important; margin-bottom: 0.35rem; font-style: normal; line-height: 1.3; font-family: 'Noto Sans Devanagari', 'Mangal', sans-serif;">साखी - कबीर</div>
  <div class="doc-header-cat-name" style="font-size: 13pt; font-weight: 600; color: #E0F2FE !important; font-style: normal; line-height: 1.3; font-family: 'Noto Sans Devanagari', 'Mangal', sans-serif;">{cat_title}</div>
</div>'''
    html_parts.append(header_html)

    # Extract Body Elements
    for elem in doc.element.body:
        if elem.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(elem, doc)
            # Skip if paragraph is duplicate header line
            p_text = p.text.strip()
            if p_text in ['साखी', 'पाठ- साखी', 'पाठ-साखी'] or ('Summary' in p_text and len(p_text) < 40) or ('Questions' in p_text and len(p_text) < 40) or ('Word Meanings' in p_text and len(p_text) < 60):
                continue
            html_parts.append(paragraph_to_html(p))
        elif elem.tag.endswith('tbl'):
            t = docx.table.Table(elem, doc)
            html_parts.append(table_to_html(t))

    # Extract Footer
    for section in doc.sections:
        if section.footer and section.footer.paragraphs:
            footer_paras = [paragraph_to_html(p, is_footer=True) for p in section.footer.paragraphs if p.text.strip()]
            if footer_paras:
                html_parts.append(
                    f'<div class="doc-footer" style="border-top: 1px solid #E2E8F0; padding-top: 1rem; margin-top: 2.5rem; text-align: right; font-style: normal; color: #156082;">'
                    + "".join(footer_paras) +
                    '</div>'
                )
            break

    return f'<div class="docx-exact-container" style="background: #ffffff; padding: 2.25rem 2rem; border-radius: 16px; box-shadow: 0 4px 20px rgba(0,0,0,0.04); font-family: \'Noto Sans Devanagari\', \'Mangal\', sans-serif; font-size: 11pt; color: #1E293B; font-style: normal;">{"".join(html_parts)}</div>'

# Convert all files
data_dir = 'D:/data'
items = os.listdir(data_dir)
target_folder = None

for item in items:
    full_path = os.path.join(data_dir, item)
    if os.path.isdir(full_path):
        subfiles = os.listdir(full_path)
        if any('saakhi' in f.lower() for f in subfiles):
            target_folder = full_path
            break

if not target_folder:
    print("Error: Target folder not found!")
    sys.exit(1)

files = os.listdir(target_folder)
converted = {}

for f in files:
    if not f.endswith('.docx') or f.startswith('~$'):
        continue
    
    fpath = os.path.join(target_folder, f)
    fn_lower = f.lower()

    if 'summary' in fn_lower:
        converted['summary'] = docx_to_exact_html(fpath, 'summary')
    elif 'notes' in fn_lower:
        converted['notes'] = docx_to_exact_html(fpath, 'notes')
    elif 'competency' in fn_lower:
        converted['competency'] = docx_to_exact_html(fpath, 'competency')
    elif 'additional' in fn_lower:
        converted['additional'] = docx_to_exact_html(fpath, 'additional')
    elif 'word_meanings' in fn_lower or 'muhavre' in fn_lower:
        converted['muhavre'] = docx_to_exact_html(fpath, 'muhavre')

json_path = 'public/chapter_html_content.json'
with open(json_path, 'r', encoding='utf-8') as jf:
    data = json.load(jf)

keys = ['cbse_10_hindi_ch1', 'cbse_10_hindi_sakhi', 'cbse_10_hindi_kabir', 'स्पर्श (भाग-2)_1', 'Sparsh_1']

for k in keys:
    data[k] = data.get(k, {})
    data[k]['title'] = 'साखी - कबीर'
    data[k]['book'] = 'स्पर्श (भाग-2)'
    data[k]['chNum'] = 1
    if 'summary' in converted:    data[k]['summary'] = converted['summary']
    if 'notes' in converted:      data[k]['notes'] = converted['notes']
    if 'competency' in converted: data[k]['competency'] = converted['competency']
    if 'additional' in converted: data[k]['additional'] = converted['additional']
    if 'muhavre' in converted:    data[k]['muhavre'] = converted['muhavre']

with open(json_path, 'w', encoding='utf-8') as jf:
    json.dump(data, jf, ensure_ascii=False, indent=2)

print("SUCCESS: CONVERTED DOCX WITH COLORED HEADER BANNER & EXACT UNIFORM FONT SIZE (18pt)!")
