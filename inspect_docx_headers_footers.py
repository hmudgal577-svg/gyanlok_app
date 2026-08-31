import os
import sys
import docx

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

data_dir = 'D:/data'
items = os.listdir(data_dir)
target_folder = None
for item in items:
    full_path = os.path.join(data_dir, item)
    if os.path.isdir(full_path):
        subfiles = os.listdir(full_path)
        if any('saakhi' in f.lower() for f in subfiles):
            target_folder = full_path
            break

print("Target folder:", target_folder)

files = os.listdir(target_folder)
for f in files:
    if not f.endswith('.docx') or f.startswith('~$'):
        continue
    fpath = os.path.join(target_folder, f)
    doc = docx.Document(fpath)
    print(f"\n=================== FILE: {f} ===================")
    for i, sec in enumerate(doc.sections):
        print(f"--- Section {i+1} ---")
        if sec.header:
            print("  [HEADER PARAGRAPHS]:")
            for hp in sec.header.paragraphs:
                print(f"    Header P: text='{hp.text}', align={hp.alignment}")
                for r in hp.runs:
                    print(f"      Run: text='{r.text}', font={r.font.name}, size={r.font.size}, bold={r.font.bold}, color={r.font.color.rgb if r.font.color else None}")
        
        if sec.footer:
            print("  [FOOTER PARAGRAPHS]:")
            for fp in sec.footer.paragraphs:
                print(f"    Footer P: text='{fp.text}', align={fp.alignment}")
                for r in fp.runs:
                    print(f"      Run: text='{r.text}', font={r.font.name}, size={r.font.size}, bold={r.font.bold}, color={r.font.color.rgb if r.font.color else None}")

